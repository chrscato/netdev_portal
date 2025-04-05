from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_imaging_template():
    doc = Document()
    
    # Title
    title = doc.add_heading('Provider Agreement', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_para.add_run('{{date}}')
    
    # Introduction
    doc.add_paragraph('This Provider Agreement (the "Agreement") is entered into as of {{date}} by and between Clarity Diagnostics, Inc. ("Clarity") and {{provider_name}} (the "Provider").')
    
    # Provider Information
    doc.add_heading('Provider Information', level=1)
    provider_info = doc.add_paragraph()
    provider_info.add_run('Provider Name: ').bold = True
    provider_info.add_run('{{provider_name}}\n')
    provider_info.add_run('DBA Name: ').bold = True
    provider_info.add_run('{{dba_name}}\n')
    provider_info.add_run('Address: ').bold = True
    provider_info.add_run('{{address}}\n')
    provider_info.add_run('Provider Type: ').bold = True
    provider_info.add_run('{{provider_type}}\n')
    provider_info.add_run('NPI: ').bold = True
    provider_info.add_run('{{npi}}\n')
    provider_info.add_run('Specialty: ').bold = True
    provider_info.add_run('{{specialty}}')
    
    # Rate Structure
    doc.add_heading('Rate Structure', level=1)
    rate_info = doc.add_paragraph()
    rate_info.add_run('Rate Type: ').bold = True
    rate_info.add_run('{{rate_type}}\n')
    rate_info.add_run('{% if rate_type == "wcfs" %}WCFS Percentage: {{wcfs_percentage}}%{% endif %}')
    
    # States
    doc.add_heading('States', level=1)
    states_para = doc.add_paragraph()
    states_para.add_run('This agreement applies to the following states: ')
    states_para.add_run('{{states}}')
    
    # Exhibit A
    doc.add_heading('Exhibit A - Rate Schedule', level=1)
    doc.add_paragraph('The following rates apply to the services provided by the Provider:')
    
    # Placeholder for the rates table
    doc.add_paragraph('{{exhibit_a}}')
    
    # Save the template
    doc.save('IMAGING_TEMPLATE.docx')
    print("IMAGING_TEMPLATE.docx created successfully!")

if __name__ == "__main__":
    create_imaging_template() 